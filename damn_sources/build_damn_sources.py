import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
)

BASE_DIR = "/home/charizard/computational-coupling/damn_sources"
MD_PATH = os.path.join(BASE_DIR, "computational_coupling_foundation.md")
DOCX_PATH = os.path.join(BASE_DIR, "computational_coupling_foundation.docx")
PDF_PATH = os.path.join(BASE_DIR, "computational_coupling_foundation.pdf")
SUMMARIES_DIR = os.path.join(BASE_DIR, "paper_summaries")
GUIDES_DIR = os.path.join(BASE_DIR, "repo_guides")

PAPERS = [
    ("01_Pais_Vieira_2013_Rat_BBI.md", "A Brain-to-Brain Interface for Real-Time Sharing of Sensorimotor Information", "Pais-Vieira et al.", "2013", "Sci Rep (10.1038/srep01319)", "Rat-to-rat motor cortex ICMS microstimulation BBI", "70%", "Invasive; binary 1-bit trigger zapping; not continuous language"),
    ("02_Rao_2014_Human_BBI.md", "A Direct Brain-to-Brain Interface in Humans", "Rao et al.", "2014", "PLoS ONE (10.1371/journal.pone.0111332)", "First non-invasive human BBI via EEG motor imagery to TMS phosphenes", "75%", "Unidirectional; low bit-rate (<1 bit/s); motor peripheral output only"),
    ("03_Foerster_2016_DIAL_MARL.md", "Learning to Communicate with Deep Multi-Agent Reinforcement Learning", "Foerster et al.", "2016", "NIPS (arxiv:1605.06676)", "Differentiable Inter-Agent Learning (DIAL) for emergent AI protocols", "75%", "Discrete grid-world toys; artificial agents without biological manifold constraints"),
    ("04_Jiang_2019_BrainNet.md", "BrainNet: A Multi-Person Brain-to-Brain Interface...", "Jiang et al.", "2019", "Sci Rep (10.1038/s41598-019-41895-7)", "3-person BBI playing Tetris via TMS phosphenes with signal reliability trust", "85%", "Bandwidth choked at discrete binary phosphenes"),
    ("05_Thual_2022_FUGW_Optimal_Transport.md", "Aligning individual brains with Fused Unbalanced Gromov-Wasserstein", "Thual et al.", "2022", "NeurIPS (arxiv:2206.09398)", "Optimal Transport whole-brain alignment matching functional signatures and cortical geometry", "70%", "Non-convex block-coordinate descent optimization cost"),
    ("06_McParlin_2022_Active_Inference.md", "Therapeutic Alliance as Active Inference", "McParlin et al.", "2022", "Front Behav Neurosci (10.3389/fnbeh.2022.897247)", "Interpersonal alignment formalization as mutual variational free energy minimization", "85%", "Purely theoretical; lacks direct neural stimulation data"),
    ("07_Jiang_2024_LaBraM_EEG_Foundation.md", "Large Brain Model for Learning Generic Representations (LaBraM)", "Jiang et al.", "2024", "ICLR (arxiv:2405.18765)", "EEG foundation model trained on 2,500 hours using VQ-NSP neural tokenizer", "80%", "Offline decoding only; lacks real-time encoding/writing"),
    ("08_Caro_2024_BrainLM_fMRI_Foundation.md", "BrainLM: A foundation model for brain activity recordings", "Caro et al.", "2024", "bioRxiv (10.1101/2023.09.12.557460)", "Generative fMRI foundation model trained on 6,700 hours across 424 AAL parcels", "70%", "Extreme BOLD latency (4-6 seconds) limits real-time interaction"),
    ("09_Scotti_2024_MindEye2.md", "MindEye2: Shared-Subject Models Enable fMRI-To-Image With 1 Hour Data", "Scotti et al.", "2024", "ICML (arxiv:2403.11207)", "Cross-subject fMRI visual perception reconstruction via shared unCLIP space", "90%", "Visual cortex specific; motion/noise sensitive"),
    ("10_Nakamura_2024_Representation_Transfer.md", "Unsupervised method for representation transfer from one brain to another", "Nakamura et al.", "2024", "Front Neuroinf (10.3389/fninf.2024.1470845)", "Zero-shot cross-subject alignment using unsupervised hyperspherical embeddings", "95% (High Threat)", "Requires static anatomical masks; passive transfer only"),
    ("11_Schreiber_2000_Transfer_Entropy.md", "Measuring Information Transfer", "Schreiber", "2000", "Phys Rev Lett (85(2):461)", "Introduced model-free, asymmetric Transfer Entropy (TE) for dynamic time series", "50%", "High sample complexity in high dimensions"),
    ("12_Shannon_1948_Information_Theory.md", "A Mathematical Theory of Communication", "Shannon", "1948", "Bell Syst Tech J (27:379)", "Foundation of information theory, entropy, and channel capacity", "Foundation", "Static discrete channels without dynamic feedback loops"),
    ("13_Hasson_2010_Neural_Coupling.md", "Speaker-listener neural coupling underlies successful communication", "Stephens & Hasson", "2010", "PNAS (107(32):14425)", "fMRI hyperscanning showing listener brain state tracks speaker brain state with lag", "65%", "Observational; cannot separate stimulus correlation from driving"),
    ("14_Chen_2015_Shared_Response_Model.md", "A Reduced-Dimension fMRI Shared Response Model", "Chen et al.", "2015", "NeurIPS (10.5555/2969442)", "Shared Response Model (SRM) factorizing multi-subject fMRI matrices", "60%", "Requires time-locked identical stimuli"),
    ("15_Montague_2002_Hyperscanning.md", "Hyperscanning: Simultaneous fMRI during linked social interactions", "Montague et al.", "2002", "NeuroImage (10.1006/nimg.2002.1150)", "Pioneered hyperscanning by linking two fMRI scanners over the internet", "50%", "High BOLD latency; dual-scanner cost")
]

REPOS = [
    ("fugw_optimal_transport_guide.md", "alexisthual/fugw", "Fused Unbalanced Gromov-Wasserstein Optimal Transport Solvers", "Pioneered by Alexis Thual & Bertrand Thirion at Inria. Used in our theory to align high-dimensional cortical surfaces based on functional similarity while penalizing topological distortion."),
    ("labram_eeg_foundation_guide.md", "935963004/LaBraM", "Large Brain Model for Learning Generic EEG Representations", "Developed by Bao-Liang Lu's group at SJTU. Pre-trained on 2,500 hours of clinical EEG using VQ-NSP. Acts as our temporal latent codec for neural state spaces."),
    ("brainlm_fmri_foundation_guide.md", "vandijklab/brainlm", "Generative fMRI Foundation Model", "Developed by David van Dijk's lab at Yale. Trained on 6,700 hours of fMRI across 424 brain parcels. Used for spatial semantic feature extraction."),
    ("biofoundation_luna_guide.md", "pulp-bio/biofoundation", "Topology-Agnostic Transformer Benchmarks", "Benchmarking suite for EEG foundation models across abnormality and artifact detection tasks."),
    ("moabb_bci_benchmark_guide.md", "NeuroDecode/MOABB", "Mother of All BCI Benchmarks", "Standardized benchmark suite ensuring reproducible evaluation across Motor Imagery, P300, and SSVEP BCI tasks.")
]

def generate_reports():
    os.makedirs(SUMMARIES_DIR, exist_ok=True)
    os.makedirs(GUIDES_DIR, exist_ok=True)
    
    for filename, title, authors, year, venue, summary, threat, weakness in PAPERS:
        filepath = os.path.join(SUMMARIES_DIR, filename)
        content = f"""# Research Paper Report: {title}

**Authors:** {authors} ({year})  
**Venue / DOI:** `{venue}`  
**Similarity Threat Level:** {threat}  

---

## 📌 Abstract & Core Summary
{summary}.

## 🎯 Main Contribution
Provides critical empirical or theoretical foundations for neural decoding, BBI, optimal transport alignment, or information theory.

## ⚠️ Critical Weakness
{weakness}.

## 🔬 Role in the Theory of Computational Coupling
Serves as an essential reference for defining Coupling Capacity C_couple, validating transfer entropy estimators, or establishing baseline performance metrics.
"""
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)

    for filename, repo_name, full_name, description in REPOS:
        filepath = os.path.join(GUIDES_DIR, filename)
        content = f"""# Open Source Repository Report: {repo_name}

**Project Name:** {full_name}  
**GitHub URL:** `https://github.com/{repo_name}`  

---

## 🚀 Overview & Role in Program
{description}

## 🛠️ Installation & Recommended Usage
```bash
git clone https://github.com/{repo_name}.git
pip install -e .
```
Used in our research pipeline to extract latent neural tokens, solve FUGW optimal transport alignment matrices, or evaluate benchmark transfer entropy models.
"""
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)

    print("All paper summaries and repo guides generated in damn_sources!")

def build_docx_and_pdf():
    with open(MD_PATH, "r", encoding="utf-8") as f:
        md_content = f.read()

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)

    lines = md_content.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped: continue
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            r = p.add_run(stripped[2:])
            r.font.bold = True
            r.font.size = Pt(18)
            r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            r = p.add_run(stripped[3:])
            r.font.bold = True
            r.font.size = Pt(14)
            r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            r = p.add_run(stripped[4:])
            r.font.bold = True
            r.font.size = Pt(12)
        elif stripped.startswith('* ') or stripped.startswith('- '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        else:
            doc.add_paragraph(stripped)

    doc.save(DOCX_PATH)
    print(f"DOCX report created at: {DOCX_PATH}")

    pdf_doc = SimpleDocTemplate(PDF_PATH, pagesize=letter, leftMargin=54, rightMargin=54, topMargin=54, bottomMargin=54)
    styles = getSampleStyleSheet()
    story = []
    
    title_style = ParagraphStyle('PDFTitle', parent=styles['Title'], fontName='Helvetica-Bold', fontSize=16, leading=20, textColor=colors.HexColor("#1F4E78"))
    h1_style = ParagraphStyle('PDFH1', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=13, leading=16, textColor=colors.HexColor("#2E75B6"))
    body_style = ParagraphStyle('PDFBody', parent=styles['Normal'], fontName='Helvetica', fontSize=9.5, leading=13.5)

    for line in lines:
        stripped = line.strip()
        if not stripped: continue
        clean_text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', stripped)
        clean_text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', clean_text)
        
        if stripped.startswith('# '):
            story.append(Paragraph(clean_text[2:], title_style))
        elif stripped.startswith('## '):
            story.append(Paragraph(clean_text[3:], h1_style))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#2E75B6"), spaceBefore=2, spaceAfter=6))
        else:
            story.append(Paragraph(clean_text, body_style))
            
    pdf_doc.build(story)
    print(f"PDF report created at: {PDF_PATH}")

if __name__ == "__main__":
    generate_reports()
    build_docx_and_pdf()
